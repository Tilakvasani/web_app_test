import os
import re
import logging
import hashlib
import httpx
from urllib.parse import unquote, quote
from typing import Dict, Any, List
from langchain_core.tools import tool, StructuredTool, BaseTool

logger = logging.getLogger("mcp_backend")

# Maximum characters for tool output before truncation (~500 tokens)
# Lower than before — smart_filter_mcp_response strips fat first so the
# truncation only kicks in as a last-resort safety net.
MAX_TOOL_OUTPUT_CHARS = 2000

def _truncate_tool_output(result: Any, tool_name: str) -> Any:
    """
    Truncates oversized tool outputs to prevent massive raw API dumps
    from inflating LLM input tokens on the ReAct re-ingestion pass.
    """
    res_str = str(result)
    original_len = len(res_str)
    if original_len > MAX_TOOL_OUTPUT_CHARS:
        truncated = res_str[:MAX_TOOL_OUTPUT_CHARS]
        logger.warning(
            f"[TOKEN SAVER] Tool '{tool_name}' output truncated from {original_len} to {MAX_TOOL_OUTPUT_CHARS} chars "
            f"(saved ~{(original_len - MAX_TOOL_OUTPUT_CHARS) // 4} tokens)."
        )
        return truncated + f"\n\n... [Output truncated from {original_len} chars. Ask for specific details if you need more.]"
    return result


import json as _json_module

def smart_filter_mcp_response(raw: Any, tool_name: str, query_hint: str = "") -> Any:
    """
    Strips fat from MCP JSON responses before they are fed back to the LLM,
    cutting token usage by 50-90% on large Notion/Cal payloads.

    Strategy:
    - If user asked for 'full content' or 'complete', skip filtering entirely.
    - For Notion pages: keep id, title, url, created/edited times, property
      *names* only (not nested objects), and a 500-char content preview.
    - For Cal bookings: keep id, title, start, end, status, attendees only.
    - For lists: process first 5 items only.
    - For all others: keep any key whose value is a scalar, truncate strings
      to 300 chars, skip large nested dicts/arrays.
    """
    # Never filter when user explicitly wants full content
    hint_lower = query_hint.lower()
    if any(kw in hint_lower for kw in ("full", "complete", "entire", "all content", "full content")):
        return raw

    def _filter_dict(d: dict, mode: str) -> dict:
        kept: dict = {}

        if mode == "notion":
            scalar_keys = {"id", "title", "url", "created_time", "last_edited_time", "type", "archived"}
            for k, v in d.items():
                if k in scalar_keys:
                    # title can be a rich-text list — flatten to plain text
                    if k == "title" and isinstance(v, list):
                        kept[k] = " ".join(t.get("plain_text", "") for t in v if isinstance(t, dict))
                    else:
                        kept[k] = v
                elif k == "properties" and isinstance(v, dict):
                    # Keep only property names, not nested value objects
                    kept["property_keys"] = list(v.keys())
                elif k in ("content", "text", "body") and isinstance(v, str):
                    kept[k] = v[:500] + ("..." if len(v) > 500 else "")
            return kept

        elif mode == "cal":
            keep_keys = {"id", "uid", "title", "startTime", "endTime", "status",
                         "attendees", "location", "description", "meetingUrl"}
            for k, v in d.items():
                if k in keep_keys:
                    if k == "description" and isinstance(v, str):
                        kept[k] = v[:200] + ("..." if len(v) > 200 else "")
                    elif k == "attendees" and isinstance(v, list):
                        kept[k] = [{"name": a.get("name"), "email": a.get("email")} for a in v[:5]]
                    else:
                        kept[k] = v
            return kept

        else:
            # Generic: keep scalars and short strings; skip large nested structures
            for k, v in d.items():
                if isinstance(v, (str, int, float, bool, type(None))):
                    kept[k] = v[:300] if isinstance(v, str) and len(v) > 300 else v
                elif isinstance(v, list) and len(v) <= 3:
                    kept[k] = v
                # skip large dicts/arrays
            return kept

    def _detect_mode(name: str) -> str:
        n = name.lower()
        if "notion" in n:
            return "notion"
        if "cal" in n or "booking" in n or "event" in n or "calendar" in n:
            return "cal"
        return "generic"

    mode = _detect_mode(tool_name)

    try:
        # If raw is a JSON string, parse it first
        parsed = raw
        if isinstance(raw, str):
            try:
                parsed = _json_module.loads(raw)
            except Exception:
                # Not JSON — truncate plain text
                return raw[:1500] + ("..." if len(raw) > 1500 else "")

        if isinstance(parsed, dict):
            filtered = _filter_dict(parsed, mode)
            result = _json_module.dumps(filtered, default=str)
        elif isinstance(parsed, list):
            filtered_list = [_filter_dict(item, mode) if isinstance(item, dict) else item
                             for item in parsed[:5]]  # max 5 items
            result = _json_module.dumps(filtered_list, default=str)
        else:
            result = str(parsed)[:1500]

        saved = len(str(raw)) - len(result)
        if saved > 0:
            logger.info(f"[TOKEN SAVER] smart_filter_mcp_response '{tool_name}': "
                        f"{len(str(raw))} → {len(result)} chars (saved ~{saved // 4} tokens).")
        return result

    except Exception as filter_err:
        logger.warning(f"[TOKEN SAVER] smart_filter_mcp_response failed for '{tool_name}': {filter_err}. Using raw.")
        return raw

def _get_tool_cache_key(tool_name: str, kwargs: dict) -> str:
    """Generates a unique cache key for a tool invocation based on name + arguments."""
    args_str = str(sorted(kwargs.items()))
    args_hash = hashlib.sha256(args_str.encode("utf-8")).hexdigest()[:32]
    # BUG FIX: use "mcp:tool_resp:*" prefix to match KEY_TOOL_RESPONSE constant in
    # redis_cache.py so cache_stats() counts them and invalidate_tool_cache() clears them
    return f"mcp:tool_resp:{tool_name}:{args_hash}"

def wrap_tool_with_coercion(tool_inst: BaseTool) -> BaseTool:
    """
    Wraps an MCP tool to intercept datatype mismatches (e.g., string instead of int)
    and coerce parameters automatically. Also applies:
    - Tool output truncation to cap token costs on massive API responses.
    - Redis response caching so repeated identical calls load instantly for 0 tokens.
    """
    async def coerced_coroutine(**kwargs):
        logger.info(f"[TOOL INVOKE] Calling tool '{tool_inst.name}' with arguments: {kwargs}")
        
        # ── Check Redis cache for identical previous call ──
        try:
            from database import cache
            cache_key = _get_tool_cache_key(tool_inst.name, kwargs)
            cached_result = await cache.get(cache_key)
            if cached_result is not None:
                logger.info(f"[TOOL CACHE HIT] Loaded cached response for '{tool_inst.name}' (0 API tokens).")
                return cached_result
        except Exception:
            cache_key = None  # Cache unavailable, proceed normally
        
        try:
            res = await tool_inst.ainvoke(kwargs)
            logger.info(f"[TOOL SUCCESS] Tool '{tool_inst.name}' successfully returned output.")
            
            # ── Smart JSON filtering (before truncation) ──
            # Strip fat from Notion/Cal/generic MCP JSON payloads so the LLM
            # only sees the fields it actually needs. Saves 50-90% of tokens
            # on large API responses without losing important data.
            res = smart_filter_mcp_response(res, tool_inst.name, str(kwargs))

            # ── Truncate oversized outputs (last-resort safety net) ──
            res = _truncate_tool_output(res, tool_inst.name)
            
            # ── Cache the result for 1 hour (MCP data is rarely stale) ──
            if cache_key:
                try:
                    from database import cache
                    await cache.set(cache_key, str(res), ttl=3600)
                except Exception:
                    pass
            
            return res
        except Exception as e:
            logger.error(f"[TOOL EXCEPTION] Tool '{tool_inst.name}' raised an error: {e}")
            err_msg = str(e)
            is_datatype_error = any(phrase in err_msg for phrase in [
                "Wrong datatype", "datatype", "type mismatch", "Parameter Input", 
                "Wrong type", "7204", "Parameter"
            ])
            if is_datatype_error:
                new_kwargs = dict(kwargs)
                coerced = False
                for param, val in new_kwargs.items():
                    if isinstance(val, str) and val.isdigit():
                        new_kwargs[param] = int(val)
                        coerced = True
                    elif isinstance(val, (int, float)):
                        new_kwargs[param] = str(int(val) if isinstance(val, float) and val.is_integer() else val)
                        coerced = True
                        
                if coerced:
                    logger.warning(
                        f"[COERCION RETRY] Datatype mismatch error detected for tool '{tool_inst.name}'. "
                        f"Attempting retry with coerced arguments. Original: {kwargs} -> Coerced: {new_kwargs}"
                    )
                    try:
                        res = await tool_inst.ainvoke(new_kwargs)
                        logger.info(f"[TOOL SUCCESS] Coerced retry of '{tool_inst.name}' completed successfully.")
                        res = smart_filter_mcp_response(res, tool_inst.name, str(new_kwargs))
                        res = _truncate_tool_output(res, tool_inst.name)
                        # BUG FIX: cache the coerced result so the same args don't repeat
                        # the error + coercion cycle on every identical call
                        if cache_key:
                            try:
                                from database import cache
                                await cache.set(cache_key, str(res), ttl=3600)
                            except Exception:
                                pass
                        return res
                    except Exception as retry_err:
                        logger.error(f"[COERCION FAIL] Coerced retry of '{tool_inst.name}' also failed: {retry_err}")
                        return f"Error: Tool '{tool_inst.name}' execution failed: {retry_err}. Please inspect this error and adjust your arguments or try another approach."
            return f"Error: Tool '{tool_inst.name}' execution failed: {e}. Please inspect this error and adjust your arguments or try another approach."

    return StructuredTool(
        name=tool_inst.name,
        description=tool_inst.description,
        args_schema=tool_inst.args_schema,
        coroutine=coerced_coroutine,
        response_format="content",
    )

# MIME type map used by both tools below
_MIME_TYPES = {
    ".pdf":  "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc":  "application/msword",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls":  "application/vnd.ms-excel",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".png":  "image/png",
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif":  "image/gif",
    ".txt":  "text/plain",
    ".md":   "text/markdown",
    ".csv":  "text/csv",
    ".json": "application/json",
    ".html": "text/html",
    ".xml":  "application/xml",
    ".zip":  "application/zip",
}

# File extensions that are plain text and can be read directly
_TEXT_EXTENSIONS = {".txt", ".md", ".json", ".csv", ".py", ".html", ".css", ".js", ".xml", ".yaml", ".yml"}

# File extensions that must be sent as binary (base64) to Google Drive / Gmail
_BINARY_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".pptx", ".png", ".jpg", ".jpeg", ".gif", ".zip"}


def get_read_uploaded_file_tool(upload_dir: str):
    """
    Creates and returns the `read_uploaded_file` tool bound to the specified upload directory.
    Use this to READ/ANALYSE file contents. For uploading files to Google Drive use
    `prepare_file_for_upload` instead.
    """
    @tool
    def read_uploaded_file(filename: str) -> str:
        """
        Read and extract the TEXT contents of a local file uploaded by the user.
        Use this tool to READ, ANALYSE, or SUMMARISE a file's content.
        For uploading a file to Google Drive or Gmail use the
        `prepare_file_for_upload` tool instead — it returns the correct
        base64-encoded binary content and MIME type that Google requires.
        """
        safe_filename = os.path.basename(filename)
        file_path = os.path.join(upload_dir, safe_filename)

        if not os.path.exists(file_path):
            return f"Error: File '{filename}' not found in the secure upload directory."

        logger.info(f"[FILE READ] Agent is reading file '{safe_filename}'...")
        ext = os.path.splitext(safe_filename)[1].lower()

        try:
            if ext in _TEXT_EXTENSIONS:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()

            elif ext == ".pdf":
                try:
                    import pypdf
                    reader = pypdf.PdfReader(file_path)
                    pages = []
                    for i, page in enumerate(reader.pages):
                        text = page.extract_text()
                        if text:
                            pages.append(f"--- Page {i+1} ---\n{text}")
                    return "\n\n".join(pages) if pages else "No text could be extracted from this PDF."
                except Exception as pdf_err:
                    return f"Error extracting text from PDF: {pdf_err}"

            elif ext == ".docx":
                try:
                    import docx
                    doc = docx.Document(file_path)
                    paragraphs = [p.text for p in doc.paragraphs if p.text]
                    tables_text = []
                    for table in doc.tables:
                        for row in table.rows:
                            tables_text.append(" | ".join(cell.text for cell in row.cells))
                    full_text = "\n".join(paragraphs)
                    if tables_text:
                        full_text += "\n\n--- Tables ---\n" + "\n".join(tables_text)
                    return full_text if full_text.strip() else "No text found in Word document."
                except Exception as doc_err:
                    return f"Error extracting text from Word document: {doc_err}"

            else:
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        return f.read(10000)
                except UnicodeDecodeError:
                    return (
                        f"File '{safe_filename}' is a binary file. "
                        f"Use `prepare_file_for_upload` to get its base64 content for uploading to Google Drive."
                    )

        except Exception as e:
            return f"Error reading file '{filename}': {e}"

    return read_uploaded_file


def get_prepare_file_for_upload_tool(upload_dir: str):
    """
    Creates and returns the `prepare_file_for_upload` tool.

    This is the CORRECT tool to use before uploading any file to Google Drive or Gmail.
    It returns:
      - base64-encoded binary content  (for PDFs, DOCX, images, etc.)
      - OR plain text content          (for .txt, .csv, .md, etc.)
      - MIME type string               (always included)

    Google Drive's `create_file` tool requires:
        name        = filename
        mimeType    = value returned by this tool
        content     = value returned by this tool  (base64 for binary, text for text)
    """
    @tool
    def prepare_file_for_upload(filename: str) -> str:
        """
        Prepare a locally uploaded file for uploading to Google Drive or Gmail.
        Returns the file's content in the correct format (base64 for binary files
        like PDF/DOCX/images, plain text for .txt/.csv/.md) and its MIME type.

        ALWAYS call this tool before calling Google Drive's create_file or
        Gmail's send_email with an attachment. Never pass extracted text from
        `read_uploaded_file` to Google Drive — that causes 'invalid document' errors
        because Google Drive expects raw binary content (base64), not extracted text.

        Returns a JSON string with keys: filename, mime_type, content, encoding.
        Pass `content` as the file content and `mime_type` as the mimeType argument
        to Google Drive's create_file tool.
        """
        import base64
        import json as _json

        safe_filename = os.path.basename(filename)
        file_path = os.path.join(upload_dir, safe_filename)

        if not os.path.exists(file_path):
            return f"Error: File '{filename}' not found in the secure upload directory."

        ext = os.path.splitext(safe_filename)[1].lower()
        mime_type = _MIME_TYPES.get(ext, "application/octet-stream")

        logger.info(f"[FILE PREP] Preparing '{safe_filename}' (MIME: {mime_type}) for upload...")

        try:
            if ext in _TEXT_EXTENSIONS:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
                encoding = "text"
            else:
                # Binary file — base64 encode it so it can be passed as a string
                # to Google Drive's create_file tool without corruption
                with open(file_path, "rb") as f:
                    raw_bytes = f.read()
                content = base64.b64encode(raw_bytes).decode("utf-8")
                encoding = "base64"

            result = {
                "filename": safe_filename,
                "mime_type": mime_type,
                "content": content,
                "encoding": encoding,
                "size_bytes": os.path.getsize(file_path),
                "instructions": (
                    f"Pass 'content' as the file content and 'mime_type' as the mimeType "
                    f"to Google Drive's create_file tool. The encoding is '{encoding}'."
                )
            }
            logger.info(f"[FILE PREP] Successfully prepared '{safe_filename}' ({encoding}, {os.path.getsize(file_path)} bytes).")
            return _json.dumps(result)

        except Exception as e:
            logger.error(f"[FILE PREP ERROR] Failed to prepare '{safe_filename}': {e}")
            return f"Error preparing file '{filename}' for upload: {e}"

    return prepare_file_for_upload

def get_query_documentation_tool(loaded_docs: Dict[str, Any]):
    """
    Creates and returns the `query_documentation` tool loaded with the active resources.
    """
    @tool
    async def query_documentation(query: str) -> str:
        """
        Search through the loaded documentation / llms.txt pages for information.
        Provide keywords or a search phrase. Returns relevant excerpts from the documents.
        """
        from database import cache, vector_search_redis

        logger.info(f"[DOC SEARCH] Initiating documentation query: '{query}' across {len(loaded_docs)} sources...")
        if not loaded_docs:
            return "No documentation resources have been loaded yet."

        # 1. Attempt Semantic Vector Search via Redis & Azure OpenAI if Redis is online
        if cache.is_available:
            logger.info(f"[DOC SEARCH] Redis active. Running semantic vector search...")
            try:
                vector_results = await vector_search_redis(query, top_k=3)
                if vector_results:
                    logger.info(f"[DOC SEARCH SUCCESS] Semantic vector search returned {len(vector_results)} matches.")
                    formatted_res = []
                    for idx, r in enumerate(vector_results):
                        formatted_res.append(
                            f"--- Match {idx+1} [Score: {r['score']:.4f}] from Doc: {r['doc_name']} ({r['url']}) ---\n"
                            f"{r['text'].strip()}"
                        )
                    return "\n\n=================================\n\n".join(formatted_res)
                else:
                    logger.info("[DOC SEARCH] Semantic vector search returned no results. Falling back to keyword search.")
            except Exception as e:
                logger.error(f"[DOC SEARCH ERROR] Semantic vector search failed: {e}. Falling back to keyword search.")

        # 2. Fallback Keyword/Regex-based Search
        logger.info("[DOC SEARCH] Running standard keyword search...")
        words = [w.lower() for w in re.findall(r"\b\w{3,}\b", query)]
        if not words:
            return "Query was too short or lacked search keywords."
        
        results = []
        for doc_name, doc in loaded_docs.items():
            content = doc["content"]
            sections = []
            current_section = []
            for line in content.split("\n"):
                if line.strip().startswith("#") and current_section:
                    sections.append("\n".join(current_section))
                    current_section = [line]
                else:
                    current_section.append(line)
            if current_section:
                sections.append("\n".join(current_section))
                
            scored_sections = []
            for sec in sections:
                sec_lower = sec.lower()
                score = sum(sec_lower.count(word) for word in words)
                if score > 0:
                    scored_sections.append((score, sec))
            
            scored_sections.sort(key=lambda x: x[0], reverse=True)
            doc_results = [sec.strip() for score, sec in scored_sections[:3]]
                
            if doc_results:
                formatted_res = f"--- From Doc: {doc_name} ({doc['url']}) ---\n" + "\n\n...\n\n".join(doc_results)
                results.append(formatted_res)
                
        if not results:
            fallback = [f"--- Doc: {doc_name} ({doc['url']}) ---\nPreview:\n{doc['content'][:600]}..." 
                        for doc_name, doc in loaded_docs.items()]
            return "No exact matches found for query. Here are previews of available docs:\n\n" + "\n\n".join(fallback)
            
        logger.info(f"[DOC SEARCH SUCCESS] Found matches for query: '{query}'")
        return "\n\n=================================\n\n".join(results)
        
    return query_documentation


@tool
def web_search(query: str) -> str:
    """
    Search the internet/Google for real-time information, news, company details,
    website URLs, phone numbers, contact links, or address details. Returns search result snippets.
    """
    # 1. Check for Tavily Key
    tavily_key = os.getenv("TAVILY_API_KEY")
    if tavily_key:
        logger.info(f"[TAVILY SEARCH] Querying Tavily: '{query}'...")
        try:
            resp = httpx.post(
                "https://api.tavily.com/search",
                json={
                    "api_key": tavily_key,
                    "query": query,
                    "search_depth": "advanced",
                    "max_results": 5
                },
                timeout=12.0
            )
            if resp.status_code == 200:
                results = resp.json().get("results", [])
                lines = []
                for idx, r in enumerate(results[:5]):
                    lines.append(f"{idx+1}. {r.get('title', 'No Title')}\n   URL: {r.get('url', '')}\n   Snippet: {r.get('content', '')}")
                if lines:
                    return "\n\n".join(lines)
        except Exception as e:
            logger.error(f"[TAVILY SEARCH ERROR] Falling back to DuckDuckGo: {e}")
            
    # 2. DuckDuckGo Fallback
    logger.info(f"[DDG SEARCH] Querying DuckDuckGo: '{query}'...")
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }
        url = f"https://html.duckduckgo.com/html/?q={quote(query)}"
        resp = httpx.get(url, headers=headers, timeout=12.0)
        if resp.status_code == 200:
            html = resp.text
            # Extract URLs and Titles
            a_tags = re.findall(r'<a\s+class="result__a"\s+href="([^"]+)">([^<]+)</a>', html)
            snippets = re.findall(r'<a\s+class="result__snippet"[^>]*>([^<]+)</a>', html)
            
            results = []
            for i in range(min(5, len(a_tags))):
                url_raw = a_tags[i][0]
                if "/l/?uddg=" in url_raw:
                    url_clean = unquote(url_raw.split("/l/?uddg=")[1].split("&")[0])
                else:
                    url_clean = url_raw
                title_clean = a_tags[i][1].strip()
                snippet_clean = snippets[i].strip() if i < len(snippets) else "No snippet available."
                
                # Clean nested tags
                title_clean = re.sub(r'<[^>]*>', '', title_clean)
                snippet_clean = re.sub(r'<[^>]*>', '', snippet_clean)
                results.append(f"{i+1}. {title_clean}\n   URL: {url_clean}\n   Snippet: {snippet_clean}")
                
            if results:
                return "\n\n".join(results)
            return "No search results found on DuckDuckGo."
    except Exception as e:
        logger.error(f"[DDG SEARCH ERROR] Failed DDG search: {e}")
        return f"Error executing web search: {e}"
    
    return "Error: Web search failed."


@tool
def web_scrape(url: str) -> str:
    """
    Scrape the webpage content from the given URL.
    Fetches the website, extracts primary text, cleans it into readable Markdown,
    and automatically saves it as a new file in your workspace files.
    """
    logger.info(f"[SCRAPE] Starting scrape for URL: {url}")
    if not url.startswith("http"):
        url = "https://" + url

    # 1. Check for Tavily Key
    tavily_key = os.getenv("TAVILY_API_KEY")
    scraped_text = ""
    
    if tavily_key:
        logger.info(f"[TAVILY EXTRACT] Querying Tavily Extract API for {url}...")
        try:
            resp = httpx.post(
                "https://api.tavily.com/extract",
                json={"api_key": tavily_key, "urls": [url]},
                timeout=15.0
            )
            if resp.status_code == 200:
                res_data = resp.json()
                results = res_data.get("results", [])
                if results and results[0].get("raw_content"):
                    scraped_text = results[0]["raw_content"]
                    logger.info("[TAVILY EXTRACT SUCCESS] Successfully extracted page content.")
        except Exception as e:
            logger.error(f"[TAVILY EXTRACT ERROR] Falling back to local/Jina: {e}")

    # 2. Direct Fetch Fallback
    if not scraped_text:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8"
        }
        try:
            resp = httpx.get(url, headers=headers, timeout=12.0, follow_redirects=True)
            if resp.status_code == 200:
                html = resp.text
                is_cf = "cloudflare" in html.lower() or "javascript is required" in html.lower() or len(html) < 200
                if not is_cf:
                    clean_html = re.sub(r'<(script|style|nav|header|footer|form)[^>]*>.*?</\1>', '', html, flags=re.DOTALL | re.IGNORECASE)
                    text = re.sub(r'<[^>]*>', ' ', clean_html)
                    text = re.sub(r'\s+', ' ', text).strip()
                    scraped_text = text[:15000]
                    logger.info("[LOCAL SCRAPE SUCCESS] Extracted clean text from HTML.")
        except Exception as e:
            logger.warning(f"[LOCAL SCRAPE FAILED] Direct fetch failed for {url}: {e}. Trying Jina proxy...")

    # 3. Jina Reader Fallback
    if not scraped_text:
        logger.info(f"[JINA SCRAPE] Querying Jina Reader Proxy for {url}...")
        try:
            jina_url = f"https://r.jina.ai/{url}"
            resp = httpx.get(jina_url, timeout=15.0)
            if resp.status_code == 200:
                scraped_text = resp.text[:15000]
                logger.info("[JINA SCRAPE SUCCESS] Successfully fetched Jina markdown.")
        except Exception as e:
            logger.error(f"[JINA SCRAPE ERROR] Jina fetch failed: {e}")
            return f"Error: Failed to scrape site '{url}'."

    if not scraped_text:
        return "Error: Could not retrieve webpage content."

    # 4. Save to workspace!
    try:
        from config import UPLOAD_DIR
        os.makedirs(UPLOAD_DIR, exist_ok=True)
        domain_match = re.search(r'https?://(?:www\.)?([^/]+)', url)
        domain = domain_match.group(1).replace(".", "_") if domain_match else "site"
        safe_filename = f"scraped_{domain}.md"
        file_path = os.path.join(UPLOAD_DIR, safe_filename)
        
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(f"# Scraped Content from: {url}\n\n{scraped_text}")
        logger.info(f"[WORKSPACE SAVE] Saved scraped page content to '{file_path}'")

        # BUG FIX: register the scraped file in state so it appears in /files and the
        # agent can reference it in future turns — previously it was saved to disk but
        # never added to state["uploaded_files"]
        try:
            import time as _time
            from database.state_manager import load_state, save_state
            state_data = load_state()
            if "uploaded_files" not in state_data:
                state_data["uploaded_files"] = []
            # Replace any stale entry with the same filename
            state_data["uploaded_files"] = [
                f for f in state_data["uploaded_files"] if f.get("name") != safe_filename
            ]
            state_data["uploaded_files"].append({
                "name": safe_filename,
                "path": os.path.abspath(file_path),
                "size": len(scraped_text.encode("utf-8")),
                "uploaded_at": _time.time()
            })
            save_state(state_data)
            logger.info(f"[WORKSPACE SAVE] Registered '{safe_filename}' in workspace state.")
        except Exception as state_err:
            logger.error(f"[WORKSPACE SAVE ERROR] Failed to register scraped file in state: {state_err}")
    except Exception as save_err:
        logger.error(f"[WORKSPACE SAVE ERROR] Failed to save scraped file: {save_err}")

    return f"Successfully scraped '{url}'! Content saved to your workspace as '{safe_filename}'.\n\nPreview of Scraped Content:\n\n{scraped_text[:1200]}..."